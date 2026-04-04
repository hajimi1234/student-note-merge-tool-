from docx import Document
from PyPDF2 import PdfMerger
import os
import sys
import datetime
import tkinter as tk
from tkinter import messagebox


def get_current_folder():
    """自动识别当前exe/脚本所在的文件夹，扔到哪就合并哪"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))


def show_message(title, content, is_error=False):
    """弹窗提示，代替控制台input，无窗口模式也能用"""
    # 隐藏主窗口
    root = tk.Tk()
    root.withdraw()
    # 弹窗置顶
    root.attributes('-topmost', True)
    if is_error:
        messagebox.showerror(title, content)
    else:
        messagebox.showinfo(title, content)
    root.destroy()


def merge_word_files(folder_path, today_str):
    output_name = f"{today_str}_笔记合并_Word版.docx"
    output_path = os.path.join(folder_path, output_name)

    # 筛选docx文件，排除临时文件和已合并文件
    word_files = [
        f for f in sorted(os.listdir(folder_path))
        if f.lower().endswith(".docx")
           and not f.startswith(f"{today_str}_笔记合并")
           and not f.startswith("~$")
    ]

    if not word_files:
        print("📄 未找到可合并的Word文件，跳过Word合并\n")
        return None

    print(f"📄 找到{len(word_files)}个Word文件，开始合并...")
    master_doc = Document()

    for file in word_files:
        file_path = os.path.join(folder_path, file)
        try:
            sub_doc = Document(file_path)
            # 100%保留原格式（字体、样式、表格）
            for para in sub_doc.paragraphs:
                new_para = master_doc.add_paragraph()
                new_para.style = para.style
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
            # 复制表格
            for table in sub_doc.tables:
                new_table = master_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
            # 每个文档后自动加分页符
            master_doc.add_page_break()
            print(f"  ✅ 已合并：{file}")
        except Exception as e:
            print(f"  ❌ 合并失败：{file}，原因：{str(e)}")

    # 保存前处理文件占用问题
    try:
        # 先删除旧文件，避免覆盖权限问题
        if os.path.exists(output_path):
            os.remove(output_path)
        master_doc.save(output_path)
        print(f"✅ Word合并完成！文件已保存：{output_name}\n")
        return output_path
    except PermissionError:
        error_msg = f"❌ Word文件 {output_name} 已被打开，请关闭后重试！"
        print(error_msg)
        show_message("错误", error_msg, is_error=True)
        return None
    except Exception as e:
        error_msg = f"❌ Word保存失败：{str(e)}"
        print(error_msg)
        show_message("错误", error_msg, is_error=True)
        return None


def merge_pdf_files(folder_path, today_str):
    output_name = f"{today_str}_笔记合并_PDF版.pdf"
    output_path = os.path.join(folder_path, output_name)

    # 筛选PDF文件
    pdf_files = [
        f for f in sorted(os.listdir(folder_path))
        if f.lower().endswith(".pdf")
           and not f.startswith(f"{today_str}_笔记合并")
    ]

    if not pdf_files:
        print("📕 未找到可合并的PDF文件，跳过PDF合并\n")
        return None

    print(f"📕 找到{len(pdf_files)}个PDF文件，开始合并...")
    merger = PdfMerger()

    for file in pdf_files:
        file_path = os.path.join(folder_path, file)
        try:
            merger.append(file_path)
            print(f"  ✅ 已合并：{file}")
        except Exception as e:
            print(f"  ❌ 合并失败：{file}，原因：{str(e)}")

    # 保存前处理文件占用问题
    try:
        # 先删除旧文件，避免覆盖权限问题
        if os.path.exists(output_path):
            os.remove(output_path)
        merger.write(output_path)
        merger.close()
        print(f"✅ PDF合并完成！文件已保存：{output_name}\n")
        return output_path
    except PermissionError:
        error_msg = f"❌ PDF文件 {output_name} 已被打开，请关闭后重试！"
        print(error_msg)
        show_message("错误", error_msg, is_error=True)
        merger.close()
        return None
    except Exception as e:
        error_msg = f"❌ PDF保存失败：{str(e)}"
        print(error_msg)
        show_message("错误", error_msg, is_error=True)
        merger.close()
        return None


if __name__ == "__main__":
    print("=" * 50)
    print("📚 学生党笔记合并神器")
    print("=" * 50)
    current_folder = get_current_folder()
    today = datetime.datetime.now().strftime("%Y%m%d")
    print(f"📂 当前处理文件夹：{current_folder}")
    print(f"📅 处理日期：{today}\n")

    # 执行合并
    word_result = merge_word_files(current_folder, today)
    pdf_result = merge_pdf_files(current_folder, today)

    # 完成提示
    print("=" * 50)
    print("🎉 所有合并任务执行完成！")
    os.startfile(current_folder)
    print("📂 已自动打开文件夹！")
    print("=" * 50)

    # 弹窗提示完成，彻底替代input()，解决lost sys.stdin
    success_msg = "✅ 合并完成！\n已自动打开文件夹，可查看合并后的文件~"
    show_message("完成", success_msg)